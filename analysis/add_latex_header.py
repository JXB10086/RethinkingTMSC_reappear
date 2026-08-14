# -*- coding: utf-8 -*-
"""给《论文复现.docx》添加与 LaTeX(ctexart) 一致的运行页眉：
- 每个一级章节独立成节，页眉显示“章节名 + 页码”（页码右对齐）；
- 标题页（第一章之前）无页眉；
- 页码连续编号；
- 页眉下方加细横线（与常见中文模板一致）。
原文件先备份为同目录 论文复现_备份.docx。
"""
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

SRC = r"F:\winoptimizeDir\Desktop\论文复现.docx"
BAK = r"F:\winoptimizeDir\Desktop\论文复现_备份.docx"
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# 从备份(未修改版)重新生成，避免旧分节结构叠加
doc = Document(BAK)
body = doc.element.body

# 1) 收集一级标题段落（styleId == "1"）
h1_paras = []
for par in doc.paragraphs:
    pPr = par._p.pPr
    if pPr is not None and pPr.pStyle is not None and pPr.pStyle.get(qn("w:val")) == "1":
        h1_paras.append(par)
print("一级标题:", [p.text for p in h1_paras])

# 2) 在每个一级标题前插入一个不可见空段落并携带分节符（下一节从新页开始）
last_sect = body.find(qn("w:sectPr"))
for par in h1_paras:
    empty = doc.add_paragraph()
    empty.paragraph_format.space_before = Pt(0)
    empty.paragraph_format.space_after = Pt(0)
    empty.paragraph_format.line_spacing = 1.0
    r = empty.add_run("")
    r.font.size = Pt(1)
    par._p.addprevious(empty._p)
    new_sect = deepcopy(last_sect)
    for tag in ("w:headerReference", "w:footerReference", "w:footnotePr", "w:endnotePr"):
        for el in new_sect.findall(qn(tag)):
            new_sect.remove(el)
    empty._p.get_or_add_pPr().append(new_sect)

# 3) 设置页眉
def set_header_run(run, text=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "宋体")
    if text is not None:
        run.text = text

def add_page_field(par):
    run = par.add_run()
    set_header_run(run)
    f1 = OxmlElement("w:fldChar")
    f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = "PAGE"
    f2 = OxmlElement("w:fldChar")
    f2.set(qn("w:fldCharType"), "end")
    run._r.append(f1)
    run._r.append(it)
    run._r.append(f2)

def add_header_rule(par):
    pPr = par._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)

sections = doc.sections
for i, par in enumerate(h1_paras, start=1):
    sec = sections[i]
    sec.header.is_linked_to_previous = False
    hp = sec.header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.tab_stops.add_tab_stop(Pt(445), WD_TAB_ALIGNMENT.RIGHT)
    r = hp.add_run(par.text)
    set_header_run(r)
    r2 = hp.add_run("\t")
    set_header_run(r2)
    add_page_field(hp)
    add_header_rule(hp)
    print(f"节 {i}: 页眉 = {par.text!r} + 页码")

# 4) 第一节（标题页）页眉留空
doc.sections[0].header.is_linked_to_previous = False
hp0 = doc.sections[0].header.paragraphs[0]
hp0.text = ""

doc.save(SRC)
print("已保存:", SRC)
print("分节数:", len(doc.sections))
