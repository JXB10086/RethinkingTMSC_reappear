# -*- coding: utf-8 -*-
"""按《论文复现.docx》排版生成完整中文小论文（原 RethinkingTMSC 转述版）。
排版：正文 NSimSun 12pt / 1.5 倍行距 / 首行缩进 2 字符 / 两端对齐；
一级标题 SimHei 16pt 居中且另起一页；二级标题 SimHei 12pt；
表格为三线表；公式居中带右对齐编号；参考文献 [n]。
"""
import shutil
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part

SRC = r"F:\winoptimizeDir\Desktop\论文复现.docx"
DST = r"F:\winoptimizeDir\Desktop\论文复现_完整版.docx"

SONG = "NSimSun"
HEI = "SimHei"
TNR = "Times New Roman"
_BMID = [1000]


def add_ref_field(p, name, display):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " REF " + name + " \\h ")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = display
    r.append(t)
    fld.append(r)
    p._p.append(fld)


def bookmark_run(run, name):
    _BMID[0] += 1
    bid = str(_BMID[0])
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), bid)
    bm_start.set(qn("w:name"), name)
    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), bid)
    run._r.addprevious(bm_start)
    run._r.addnext(bm_end)


REF_TOK = re.compile(
    r"~([^~]+)~|\^([^^]+)\^|§T([1-4])§|§E([1-7])§|\[[0-9]+(?:[,-][0-9]+)*\]")


def emit_text(p, text, font=SONG, size=12):
    pos = 0
    for m in REF_TOK.finditer(text):
        if m.start() > pos:
            set_run(p.add_run(), text[pos:m.start()], font=font, size=size)
        tok = m.group(0)
        if tok.startswith("§T"):
            add_ref_field(p, "tbl" + m.group(3), "表 " + m.group(3))
        elif tok.startswith("§E"):
            add_ref_field(p, "eq" + m.group(4), "(" + m.group(4) + ")")
        elif m.group(1) is not None:
            r = p.add_run()
            set_run(r, m.group(1), font=font, size=size)
            r.font.subscript = True
        elif m.group(2) is not None:
            r = p.add_run()
            set_run(r, m.group(2), font=font, size=size)
            r.font.superscript = True
        else:
            r = p.add_run()
            set_run(r, tok, font=font, size=size)
            first = int(re.search(r"\d+", tok).group())
            hl = OxmlElement("w:hyperlink")
            hl.set(qn("w:anchor"), "ref" + str(first))
            hl.set(qn("w:history"), "1")
            r._r.addprevious(hl)
            hl.append(r._r)
        pos = m.end()
    if pos < len(text):
        set_run(p.add_run(), text[pos:], font=font, size=size)


def set_run(run, text=None, font=SONG, size=12, bold=False, italic=False, color=None):
    if text is not None:
        run.text = text
    run.font.name = TNR
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), TNR if font in (SONG, HEI) else font)
    rFonts.set(qn("w:hAnsi"), TNR if font in (SONG, HEI) else font)
    rFonts.set(qn("w:eastAsia"), font)


def body(doc, text, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12,
         font=SONG, bold_prefix=None, space_after=0, line=1.5):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = line
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    if indent:
        pf.first_line_indent = Pt(size * 2)
    if bold_prefix:
        set_run(p.add_run(), bold_prefix, font=font, size=size, bold=True)
    emit_text(p, text, font=font, size=size)
    return p


def h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.page_break_before = True
    pf.line_spacing = 1.329
    pf.space_before = Pt(0)
    set_run(p.add_run(), text, font=HEI, size=16)
    return p


def h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(1.5)
    pf.space_after = Pt(1.5)
    set_run(p.add_run(), text, font=HEI, size=12)
    return p


def formula(doc, segments, num):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.2
    pf.tab_stops.add_tab_stop(Cm(15.9), WD_ALIGN_PARAGRAPH.RIGHT)
    for kind, text in segments:
        r = p.add_run()
        set_run(r, text, font=TNR, size=12, italic=True)
        if kind == "sub":
            r.font.subscript = True
        elif kind == "sup":
            r.font.superscript = True
    set_run(p.add_run(), "\t", font=TNR, size=12)
    nr = p.add_run()
    set_run(nr, "(" + num + ")", font=TNR, size=12)
    bookmark_run(nr, "eq" + num)
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.2
    m = re.match(r"^(表\s*(\d+))(.*)$", text)
    if m:
        r1 = p.add_run()
        set_run(r1, m.group(1), font=HEI, size=10.5)
        bookmark_run(r1, "tbl" + m.group(2))
        set_run(p.add_run(), m.group(3), font=HEI, size=10.5)
    else:
        set_run(p.add_run(), text, font=HEI, size=10.5)
    return p


def _set_tbl_borders(tbl, top_sz=12, bottom_sz=12, header_sz=6):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for tag, sz in (("top", top_sz), ("bottom", bottom_sz)):
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), "000000")
        borders.append(el)
    for tag in ("left", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    tblPr.append(borders)


def _header_bottom_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)
    tcPr.append(borders)


def add_table(doc, header, rows, widths_cm, caption_text, font_size=9):
    caption(doc, caption_text)
    tbl = doc.add_table(rows=len(rows) + 1, cols=len(header))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    _set_tbl_borders(tbl)
    for j, h in enumerate(header):
        cell = tbl.rows[0].cells[j]
        cell.width = Cm(widths_cm[j])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.15
        set_run(p.add_run(), h, font=SONG, size=font_size)
        _header_bottom_border(cell)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = tbl.rows[i].cells[j]
            cell.width = Cm(widths_cm[j])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.15
            set_run(p.add_run(), str(val), font=SONG, size=font_size)
    body(doc, "", indent=False, space_after=4)
    return tbl


def build():
    global DST
    try:
        shutil.copy2(SRC, DST)
    except PermissionError:
        DST = DST.replace(".docx", "v2.docx")
        print("原文件被占用，另存为:", DST)
        shutil.copy2(SRC, DST)
    doc = Document(DST)
    body_el = doc.element.body
    # 清空正文（保留 sectPr）
    for child in list(body_el):
        if child.tag != qn("w:sectPr"):
            body_el.remove(child)

    # ============ 标题区 ============
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(), "目标导向多模态情感分类的实证研究", font=HEI, size=18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    set_run(p.add_run(), "——基于 RethinkingTMSC（EMNLP 2023 Findings）的中文转述", size=12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_run(p.add_run(), "简相彪　　2026 年 8 月", size=10.5)

    body(doc,
         "近年来，目标导向多模态情感分类（Target-oriented Multimodal Sentiment "
         "Classification，TMSC）受到学者们的广泛关注，但现有多模态模型的性能已逐渐进入"
         "瓶颈。为探究其成因，本文对主流模型结构与两个常用数据集开展系统的实证研究，围绕"
         "三个问题展开：（1）各模态对 TMSC 是否同等重要？（2）哪类多模态融合模块更有效？"
         "（3）现有数据集能否充分支撑该任务的研究？实验与数据分析表明：现有 TMSC 系统主要"
         "依赖文本模态，多数目标的情感仅凭文本即可判定；各类融合模块相对纯文本模型没有显著"
         "增益，部分融合方式甚至导致性能下降；数据集在规模、标注质量以及“图文共同决定情感”"
         "的样本占比等方面均存在明显局限。基于上述发现，本文从模型设计与数据集构建两个角度"
         "给出若干建议，以期为 TMSC 的未来研究提供参考。",
         indent=False, bold_prefix="摘要：")
    body(doc,
         "目标导向多模态情感分类；多模态融合；BERT；视觉编码器；实证研究",
         indent=False, bold_prefix="关键词：")

    # ============ 1 引言 ============
    h1(doc, "1引言")
    body(doc,
         "目标导向情感分类又称方面级情感分类，是情感分析中的基础任务，其目标是判断文本中"
         "指定目标实体所表达的情感极性（正面、中性或负面）[11]。为了进一步利用推文中的图像"
         "信息，Yu 和 Jiang[2] 提出目标导向多模态情感分类（TMSC），将文本与图像两种模态的"
         "信息相结合以判断目标的情感。")
    body(doc,
         "近年来，TMSC 系统的性能逐渐趋于平台期：在广泛使用的 Twitter15 与 Twitter17 两个"
         "数据集上，现有最优基线模型的 Macro-F1 仅约 70[1]。为弄清性能停滞的原因，有必要"
         "从模型层面与模态层面分别进行分析。粗略来看，现有模型结构可归结为两类模块：一是"
         "对各模态进行表示学习的单模态编码器，二是建模模态间交互的多模态融合模块[1]。")
    body(doc,
         "本文在统一实验设置下系统比较不同编码器与融合模块，并对两个数据集的特性进行深入"
         "分析，旨在回答以下三个问题：")
    body(doc, "Q1：各模态对 TMSC 是否同等重要？为回答该问题，本文分别评测文本与图像单模态"
              "模型：文本侧采用预测练语言模型 BERT[3]，图像侧采用 ResNet-152[4]、ViT[5] 与 "
              "Faster R-CNN[6] 三种主流视觉编码器。")
    body(doc, "Q2：哪类多模态融合模块更有效？本文将现有融合策略归纳为拼接（Concatenation）、"
              "张量融合（Tensor Fusion）[7]、自注意力（Self Attention）、图像到文本（Image2Text）、"
              "文本到图像（Text2Image）与双向（Bi-direction）六类，并在统一设置下进行对比。")
    body(doc, "Q3：现有数据集能否充分支撑该任务的研究？本文对 Twitter15 与 Twitter17 进行"
              "深入分析，得到四点发现：（1）数据集规模有限；（2）多模态情感与文本情感的一致性"
              "远高于其与视觉情感的一致性；（3）大量目标并未出现在图像中；（4）仅少数样本的"
              "情感需要文本与图像共同决定[1]。")
    body(doc,
         "本文的主要贡献包括：（1）系统考察了不同单模态编码器与多模态融合模块对 TMSC 性能"
         "的影响；（2）深入分析了现有数据集的不足；（3）基于实验观察给出若干研究建议，为 "
         "TMSC 的模型设计与数据集构建提供参考[1]。")

    # ============ 2 相关工作 ============
    h1(doc, "2相关工作")
    body(doc,
         "作为情感分析的子任务，TMSC 近年来受到广泛关注[2]。Xu 等[19]构建了中文数据集 "
         "Multi-ZOL，并提出多跳记忆网络以建模模态间交互；随后 Yu 和 Jiang[2] 构建了 Twitter15 "
         "与 Twitter17 两个英文数据集，并将 BERT 引入该任务。后续研究大致沿两个方向展开："
         "一是不断探索增强模态间交互的方法[14]，二是将预测练模型应用于该任务[15,16]。尽管"
         "已有大量工作，多模态模型相对纯文本模型仍未取得显著增益，这正是本文开展实证研究"
         "的出发点。在文本建模方面，预测练语言模型 BERT[3] 已取代早期基于 LSTM[12] 与记忆"
         "网络[13] 的方法，成为该任务的主流文本编码器。")

    # ============ 3 任务定义与模型体系 ============
    h1(doc, "3任务定义与模型体系")
    h2(doc, "3.1任务定义")
    body(doc,
         "给定一条推文文本 X、文本中指定的目标实体 t 以及推文配图 I，TMSC 的目标是判断目标 "
         "t 被表达的情感极性 y∈{正面，中性，负面}。评测指标为准确率（ACC）与宏平均 F1"
         "（Macro-F1）[1]。")
    h2(doc, "3.2单模态编码器")
    body(doc,
         "文本编码器采用 BERT[3]。对于图像模态，本文考察三种主流编码器：")
    body(doc,
         "（1）ResNet-152[4]：将图像缩放为 224×224 后输入网络，取最后一层卷积输出的 49 个"
         "区域作为图像表示，如公式§E1§ 所示：")
    formula(doc, [("t", "I = [v"), ("sub", "1"), ("t", ", v"), ("sub", "2"),
                  ("t", ", …, v"), ("sub", "49"), ("t", "], v"), ("sub", "i"),
                  ("t", " ∈ R"), ("sup", "2048")], "1")
    body(doc, "其中 v~i~ 表示第 i 个区域的 2048 维特征向量。")
    body(doc,
         "（2）ViT[5]：将图像划分为 16×16 的 patch，并在序列前端加入分类标记（CLS token），"
         "经 Transformer[8] 编码后得到图像表示，如公式§E2§ 所示：")
    formula(doc, [("t", "I = [v"), ("sub", "cls"), ("t", ", v"), ("sub", "1"),
                  ("t", ", v"), ("sub", "2"), ("t", ", …, v"), ("sub", "196"),
                  ("t", "], v"), ("sub", "i"), ("t", " ∈ R"), ("sup", "768")], "2")
    body(doc,
         "（3）Faster R-CNN[6]：采用在 Visual Genome[17] 上重新训练的目标检测模型，选取置信度"
         "最高的 36 个目标提议（object proposals）作为图像表示，如公式§E3§ 所示，")
    formula(doc, [("t", "I = [v"), ("sub", "1"), ("t", ", v"), ("sub", "2"),
                  ("t", ", …, v"), ("sub", "36"), ("t", "], v"), ("sub", "i"),
                  ("t", " ∈ R"), ("sup", "2048")], "3")
    body(doc,
         "其中 v~i~ 来自区域提议网络（RPN）的 ROI 池化层。为统一不同图像编码器的输出维度，"
         "融合前通过一个线性映射层将图像表示映射到 768 维[1]。")
    h2(doc, "3.3多模态融合模块")
    body(doc,
         "设池化后的文本表示为 H~p~^T^∈R^768^，池化后的图像表示为 H~p~^I^∈R^768^，本文按以下六类融合"
         "策略进行对比[1]：")
    body(doc,
         "（1）拼接（Concatenate）：将文本与图像池化表示直接拼接，得到多模态表示"
         "（如公式§E4§ 所示）")
    formula(doc, [("t", "H = H"), ("sub", "p"), ("sup", "I"), ("t", " ⊕ H"),
                  ("sub", "p"), ("sup", "T"), ("t", ",  H ∈ R"), ("sup", "768+768")], "4")
    body(doc,
         "（2）张量融合（Tensor Fusion）[7]：通过外积建模模态间的交互，同时保留各模态自身的"
         "特性（如公式§E5§ 所示），")
    formula(doc, [("t", "H = H"), ("sub", "p"), ("sup", "I"), ("t", " ⊗ H"),
                  ("sub", "p"), ("sup", "T"), ("t", ",  H ∈ R"), ("sup", "768×768")], "5")
    body(doc,
         "（3）自注意力（Self Attention）：将图像表示与文本表示拼接后，经过三层自注意力层"
         "与池化层，得到（如公式§E6§ 所示）")
    formula(doc, [("t", "H = SelfAttn([H"), ("sup", "T"), ("t", "; H"),
                  ("sup", "I"), ("t", "]) ∈ R"), ("sup", "768")], "6")
    body(doc,
         "（4）图像到文本（Image2Text）与文本到图像（Text2Image）：前者以图像表示作为查询"
         "（Query）、文本表示作为键值（Key/Value），经三层交叉注意力得到融合表示；后者将两者"
         "角色互换。将二者结果拼接，即得到双向（Bi-direction）表示（如公式§E7§ 所示）[1]：")
    formula(doc, [("t", "H = [H"), ("sup", "I→T"), ("t", "; H"), ("sup", "T→I"),
                  ("t", "] ∈ R"), ("sup", "768+768")], "7")
    body(doc,
         "其中 H^I→T^ 与 H^T→I^ 分别表示 Image2Text 与 Text2Image 的输出。§T1§ 汇总了本文涉及的"
         "各类基线模型的结构差异。")

    add_table(
        doc,
        ["模型", "文本编码器", "图像编码器", "融合策略"],
        [
            ["mBERT[2]", "BERT", "ResNet-152", "拼接"],
            ["TomBERT[2]", "BERT", "ResNet-152", "拼接 + 图像到文本"],
            ["EF-CapTrBERT[14]", "BERT", "ResNet-152", "跨模态交互"],
            ["Res-BERT+BL[1]", "BERT", "ResNet-152", "自注意力"],
            ["Res-BERT+BL-TFN[1]", "BERT", "ResNet-152", "张量融合"],
            ["SMP[15]", "BERT", "ViT", "双向注意力"],
            ["VLP[16]", "预测练视觉语言模型", "Faster R-CNN", "预测练"],
        ],
        [3.0, 3.4, 3.2, 6.3],
        "表 1　各类基准模型的结构总览[1]",
    )

    # ============ 4 实验设置与结果分析 ============
    h1(doc, "4实验设置与结果分析")
    h2(doc, "4.1实验设置")
    body(doc,
         "实验在 Twitter15 与 Twitter17 两个英文数据集上进行[2]。统一超参数设置如下：训练 8 "
         "个 epoch，批大小 32，学习率 2×10^-5^，优化器采用 Adam[18]，并在 5 个随机种子"
         "（0/42/199/2022/11122）上重复实验，报告均值与标准差[1]。")
    h2(doc, "4.2总体结果")
    body(doc,
         "§T2§ 给出各单模态模型与多模态模型在两个数据集上的实验结果。整体来看，可以得出"
         "以下四点观察[1]：")
    body(doc,
         "第一，纯文本模型（BERT）表现良好，而纯视觉模型（ResNet、ViT、Faster R-CNN）表现"
         "较差，说明在该任务上模型对文本的依赖远大于图像，且这一现象在 Twitter15 上更为"
         "明显。")
    body(doc,
         "第二，融合方式对性能有显著影响：以获取文本信息为主的融合（如 Image2Text）优于以"
         "获取图像信息为主的融合（如 Text2Image），再次印证文本与图像在任务中的重要性并不"
         "对等。")
    body(doc,
         "第三，与纯文本模型相比，各类融合模块并未带来显著增益，部分融合甚至更差。其原因"
         "在于部分图像并未提供相关信息，反而引入了干扰信息。")
    body(doc,
         "第四，图像编码器的影响并不明确：单模态图像模型性能低且方差大，而在多模态融合"
         "设置下，不同图像编码器之间的性能差异很小。这与现有数据集中视觉数据的特性有关，"
         "详见第 5 节的数据分析。")
    body(doc,
         "基于上述分析，未来 TMSC 模型设计应重点关注：（1）充分利用文本信息的优势；"
         "（2）设计更有效的图像编码方法，以更好地提取图像语义信息；（3）增强融合模块的"
         "抗噪能力，使其能够灵活选择并利用文本与图像中的有效特征[1]。")

    add_table(
        doc,
        ["模态", "模型", "融合", "T15 ACC", "T15 F1", "T17 ACC", "T17 F1"],
        [
            ["文本", "BERT", "—", "76.72±1.16", "71.19±2.19", "68.04±0.40", "65.66±0.35"],
            ["图像", "ResNet", "—", "57.65±1.00", "32.52±2.66", "57.79±0.99", "51.98±1.23"],
            ["图像", "ViT", "—", "59.65±1.13", "31.25±2.71", "59.53±0.95", "54.08±0.78"],
            ["图像", "Faster R-CNN", "—", "55.97±1.10", "35.72±5.43", "56.18±0.85", "49.88±1.70"],
            ["多模态", "ResNet", "拼接", "75.29±0.45", "68.71±1.34", "67.92±0.56", "65.32±0.53"],
            ["多模态", "ResNet", "张量融合", "74.19±0.94", "68.93±0.57", "66.66±1.21", "63.99±1.61"],
            ["多模态", "ResNet", "自注意力", "76.03±0.96", "70.57±2.39", "68.01±0.96", "65.41±1.60"],
            ["多模态", "ResNet", "Image2Text", "77.13±1.33", "71.48±1.90", "69.37±0.36", "66.85±0.79"],
            ["多模态", "ResNet", "Text2Image", "75.18±1.66", "67.77±4.81", "68.07±0.58", "65.18±1.48"],
            ["多模态", "ResNet", "双向", "77.32±0.63", "72.06±0.81", "68.41±1.01", "66.39±1.39"],
            ["多模态", "ViT", "拼接", "76.22±0.90", "70.37±1.45", "67.94±0.70", "66.17±0.78"],
            ["多模态", "ViT", "张量融合", "73.44±0.78", "67.46±1.45", "65.46±1.67", "62.02±1.40"],
            ["多模态", "ViT", "自注意力", "75.08±0.41", "68.94±0.83", "67.52±0.58", "65.56±0.35"],
            ["多模态", "ViT", "Image2Text", "77.11±0.44", "71.91±0.42", "69.14±0.52", "66.96±0.68"],
            ["多模态", "ViT", "Text2Image", "75.12±1.01", "69.40±1.38", "67.52±1.06", "64.49±1.46"],
            ["多模态", "ViT", "双向", "76.70±0.75", "71.67±1.45", "69.16±0.17", "67.25±0.56"],
            ["多模态", "Faster", "拼接", "75.45±0.73", "69.77±1.23", "67.60±1.15", "64.74±1.69"],
            ["多模态", "Faster", "张量融合", "72.09±0.66", "66.77±1.04", "66.34±1.45", "62.96±2.09"],
            ["多模态", "Faster", "自注意力", "76.09±0.89", "70.08±1.37", "68.09±1.10", "66.12±1.23"],
            ["多模态", "Faster", "Image2Text", "77.36±0.37", "71.69±0.37", "68.43±0.65", "66.44±1.10"],
            ["多模态", "Faster", "Text2Image", "70.82±2.99", "57.94±5.81", "60.31±6.43", "54.50±7.06"],
            ["多模态", "Faster", "双向", "76.57±0.46", "70.88±0.89", "69.51±0.62", "67.50±0.37"],
        ],
        [1.5, 2.2, 2.1, 2.6, 2.6, 2.6, 2.6],
        "表 2　Twitter15 与 Twitter17 上的实验结果（5 种子均值±标准差）[1]",
        font_size=8,
    )

    # ============ 5 数据集分析 ============
    h1(doc, "5数据集分析")
    h2(doc, "5.1数据集统计")
    body(doc,
         "§T3§ 给出两个数据集的统计信息。可以看出：（1）数据集规模较小，每个样本平均包含的"
         "目标数不足 1.5 个；（2）情感标签分布不均衡，中性样本约占 50%，负面样本不足 15%。"
         "造成上述现象的原因在于，Twitter15 与 Twitter17 最初分别是为命名实体识别任务构建"
         "的[9,10]，而非专门为 TMSC 设计[1]。")
    add_table(
        doc,
        ["划分", "T15 负面", "T15 中性", "T15 正面", "T15 总数", "T15 均目标",
         "T17 负面", "T17 中性", "T17 正面", "T17 总数", "T17 均目标"],
        [
            ["训练", "368", "1883", "928", "3179", "1.348",
             "416", "1638", "1508", "3562", "1.410"],
            ["开发", "149", "670", "303", "1122", "1.336",
             "144", "517", "515", "1176", "1.439"],
            ["测试", "113", "607", "317", "1037", "1.354",
             "168", "573", "493", "1234", "1.450"],
        ],
        [1.3, 1.6, 1.6, 1.6, 1.5, 1.6, 1.6, 1.6, 1.6, 1.5, 1.6],
        "表 3　两个数据集的统计信息（“均目标”指每个样本的平均目标数）[1]",
        font_size=8,
    )
    h2(doc, "5.2标注层面分析")
    body(doc,
         "为进一步理解性能瓶颈，本文参照 Yu 和 Jiang[2] 的标注流程，邀请三位领域专家从"
         "Twitter15 与 Twitter17 测试集中各随机抽取 200 条样本，从文本情感一致性、视觉情感"
         "一致性、目标是否存在以及情感是否由图文共同决定四个方面进行标注，并采用多数投票"
         "作为最终标注结果（见§T4§）[1]。分析得到以下发现：")
    body(doc,
         "第一，多模态情感与文本情感具有高度一致性，但与视觉情感的一致性较低。在 Twitter15 "
         "上，93% 的目标其文本情感与多模态情感一致，而视觉情感一致的比例仅为 47.5%。这说明"
         "数据集存在明显的分布偏差，即文本信息对判定多模态情感更为有效。该现象在 Twitter17 "
         "上有所缓解，但文本信息的一致性仍高于视觉信息。")
    body(doc,
         "第二，大量目标并未出现在图像中，这不符合目标导向多模态情感分类任务的设定。其原因"
         "可能在于数据集构建时目标直接从文本中选取，并未考虑对应图像的内容[2]。")
    body(doc,
         "第三，由于图像与目标不相关以及目标在图像中缺失，仅有少量样本的情感需要文本与图像"
         "共同决定：Twitter15 上约为 22%，Twitter17 上约为 55%。就多模态任务而言，这两个"
         "数据集可能并非最合适的评测基准[1]。")
    add_table(
        doc,
        ["标注指标", "Twitter15", "Twitter17"],
        [
            ["文本情感与多模态情感一致", "93%", "较高（仍高于视觉一致性）"],
            ["视觉情感与多模态情感一致", "47.5%", "较低（较 Twitter15 有所缓解）"],
            ["情感需图文共同决定", "22%", "55%"],
        ],
        [5.5, 4.5, 5.9],
        "表 4　400 条随机样本的人工标注分析（三位专家多数投票）[1]",
    )
    body(doc,
         "基于上述分析，本文认为高质量 TMSC 数据集应具备以下特征：（1）真实反映现实世界"
         "的数据分布，包括标签不均衡等情况，并为不同情形提供充足的样本；（2）具有较大的"
         "数据多样性，涵盖多种数据类型与领域，以有效检验模型的泛化能力与鲁棒性；"
         "（3）提供多维度标注信息，包括多模态情感与单模态情感，以支持对模型处理不同数据"
         "来源能力的深入分析[1]。")

    # ============ 6 结论与展望 ============
    h1(doc, "6结论与展望")
    body(doc,
         "本文对 TMSC 任务开展了系统的实证研究与数据分析。结果表明：现有多模态模型相对"
         "纯文本模型并未取得显著性能提升，其主要原因在于现有数据集对文本模态的过度依赖，"
         "视觉模态发挥的作用相对有限。基于实验分析，本文从模型设计与数据集构建两个方面"
         "提出了未来研究方向，以更好地体现社交媒体情感的多模态特性。")
    body(doc,
         "本文工作仍存在一定局限：（1）数据分析主要针对当前公开的英文数据集 Twitter15 与 "
         "Twitter17，未涉及研究较少的中文数据集 Multi-ZOL[19]；（2）虽然分析指出了现有数据"
         "集的问题，但并未构建新的、更合适的数据集；（3）实验未专门比较不同文本编码方法的"
         "影响，而将研究重点放在图像编码方法与融合模块上[1]。上述问题将作为后续工作继续"
         "探索。")

    # ============ 参考文献 ============
    h1(doc, "参考文献")
    refs = [
        "[1] Ye J, Zhou J, Tian J, et al. RethinkingTMSC: An Empirical Study for Target-Oriented "
        "Multimodal Sentiment Classification[C]. Findings of ACL: EMNLP 2023: 270–277.",
        "[2] Yu J, Jiang J. Adapting BERT for Target-Oriented Multimodal Sentiment "
        "Classification[C]. IJCAI 2019: 5408–5414.",
        "[3] Devlin J, Chang M W, Lee K, et al. BERT: Pre-training of Deep Bidirectional "
        "Transformers for Language Understanding[C]. NAACL-HLT 2019: 4171–4186.",
        "[4] He K, Zhang X, Ren S, et al. Deep Residual Learning for Image Recognition[C]. "
        "CVPR 2016: 770–778.",
        "[5] Dosovitskiy A, Beyer L, Kolesnikov A, et al. An Image is Worth 16×16 Words: "
        "Transformers for Image Recognition at Scale[C]. ICLR 2021.",
        "[6] Ren S, He K, Girshick R, et al. Faster R-CNN: Towards Real-Time Object Detection "
        "with Region Proposal Networks[C]. NeurIPS 2015: 91–99.",
        "[7] Zadeh A, Chen M, Poria S, et al. Tensor Fusion Network for Multimodal Sentiment "
        "Analysis[C]. EMNLP 2017: 1103–1114.",
        "[8] Vaswani A, Shazeer N, Parmar N, et al. Attention is All You Need[C]. NeurIPS 2017: "
        "5998–6008.",
        "[9] Zhang Q, Fu J, Liu X, et al. Adaptive Co-attention Network for Named Entity "
        "Recognition in Tweets[C]. AAAI 2018: 5674–5681.",
        "[10] Lu D, Neves L, Carvalho V, et al. Visual Attention Model for Name Tagging in "
        "Multimodal Social Media[C]. ACL 2018: 1990–1999.",
        "[11] Pontiki M, Galanis D, Pavlopoulos J, et al. SemEval-2014 Task 4: Aspect Based "
        "Sentiment Analysis[C]. SemEval@COLING 2014: 27–35.",
        "[12] Hochreiter S, Schmidhuber J. Long Short-Term Memory[J]. Neural Computation, "
        "1997, 9(8): 1735–1780.",
        "[13] Weston J, Chopra S, Bordes A. Memory Networks[C]. ICLR 2015.",
        "[14] Khan Z, Fu Y. Exploiting BERT for Multimodal Target Sentiment Classification "
        "through Input Space Translation[C]. ACM MM 2021: 3034–3042.",
        "[15] Ye J, Zhou J, Tian J, et al. Sentiment-Aware Multimodal Pre-training for "
        "Multimodal Sentiment Analysis[J]. Knowledge-Based Systems, 2022, 258: 110021.",
        "[16] Ling Y, Yu J, Xia R. Vision-Language Pre-training for Multimodal Aspect-Based "
        "Sentiment Analysis[C]. ACL 2022: 2149–2159.",
        "[17] Krishna R, Zhu Y, Groth O, et al. Visual Genome: Connecting Language and Vision "
        "Using Crowdsourced Dense Image Annotations[J]. IJCV, 2017, 123(1): 32–73.",
        "[18] Kingma D P, Ba J. Adam: A Method for Stochastic Optimization[C]. ICLR 2015.",
        "[19] Xu N, Mao W, Chen G. Multi-interactive Memory Network for Aspect Based Multimodal "
        "Sentiment Analysis[C]. AAAI 2019: 371–378.",
    ]
    for idx, ref in enumerate(refs, start=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        pf.space_after = Pt(0)
        pf.left_indent = Pt(24)
        pf.first_line_indent = Pt(-24)
        text = ref
        first_run = p.add_run()
        set_run(first_run, text, font=SONG, size=10.5)
        # 书签 refN：供文内引用跳转
        _BMID[0] += 1
        bid = str(_BMID[0])
        bs = OxmlElement("w:bookmarkStart")
        bs.set(qn("w:id"), bid)
        bs.set(qn("w:name"), "ref" + str(idx))
        be = OxmlElement("w:bookmarkEnd")
        be.set(qn("w:id"), bid)
        first_run._r.addprevious(bs)
        first_run._r.addnext(be)

    # 打开文档时自动更新交叉引用域
    settings = doc.settings.element
    uf = OxmlElement("w:updateFields")
    uf.set(qn("w:val"), "true")
    settings.append(uf)

    doc.save(DST)
    print("saved:", DST)


if __name__ == "__main__":
    build()
