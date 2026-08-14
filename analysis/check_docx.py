# -*- coding: utf-8 -*-
"""DOCX 结构检查：标题/表格/图片/伪项目符号/样式统计。"""
from collections import Counter
from docx import Document

path = r"F:\winoptimizeDir\Desktop\RethinkingTMSC\论文\中文实证论文.docx"
doc = Document(path)

heads = [(par.style.name, par.text) for par in doc.paragraphs
         if par.style.name.startswith("Heading")]
print("headings:", len(heads))
for s, t in heads:
    print(" ", s, "|", t)

print("tables:", len(doc.tables))
imgs = [r for r in doc.part.related_parts.values() if "image" in r.content_type]
print("images:", len(imgs))

fake = [par.text[:40] for par in doc.paragraphs
        if par.text.strip().startswith(("\u2022", "- ")) and
        par.style.name not in ("List Bullet", "List Number")]
print("fake bullets:", len(fake))

c = Counter(par.style.name for par in doc.paragraphs)
print("style counts:", dict(c))

alltext = "\n".join(p.text for p in doc.paragraphs)
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            alltext += "\n" + cell.text
for key in ["摘要", "1　引言", "2　任务定义", "3　数据集分析", "4　复现设置",
            "5　结果与实证分析", "6　复现性讨论", "7　结论与展望", "参考文献",
            "附录 A", "表 1", "表 2", "表 3", "表 4", "表 5", "表 6",
            "图 1", "图 2", "图 3", "图 4", "图 5", "图 6",
            "Res22Bert", "Faster2Bert", "RethinkingTMSC"]:
    print(("OK  " if key in alltext else "MISS"), key)

# 页面与样式设置检查
s = doc.sections[0]
print("page twips:", s.page_width.twips, "x", s.page_height.twips)
print("margins in:", s.top_margin.inches, s.bottom_margin.inches,
      s.left_margin.inches, s.right_margin.inches)
h1 = doc.styles["Heading 1"]
print("H1 page_break_before:", h1.paragraph_format.page_break_before)
print("H1 size/color:", h1.font.size.pt, h1.font.color.rgb)
n = doc.styles["Normal"]
ea = n.element.rPr.rFonts.get(
    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia")
print("Normal size/eastAsia:", n.font.size.pt, ea)
for par in doc.paragraphs:
    if par.style.name == "Heading 1":
        print("H1 pagebreak:", par.paragraph_format.page_break_before, "|", par.text)
