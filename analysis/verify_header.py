# -*- coding: utf-8 -*-
"""验证修改后的 docx：分节数、各节页眉文本、各节首页段落。"""
from docx import Document

p = r"F:\winoptimizeDir\Desktop\论文复现.docx"
doc = Document(p)
secs = doc.sections
print("分节数:", len(secs))
for i, sec in enumerate(secs):
    htext = "".join(r.text for r in sec.header.paragraphs[0].runs)
    print(f"节{i}: header={htext!r} linked={sec.header.is_linked_to_previous}")

# 各节第一段（确认章节分节正确）
for i, sec in enumerate(secs):
    first = None
    for par in doc.paragraphs:
        # 粗略：按顺序取分节后的首个非空段落较复杂，这里只打印全部一级标题位置
        pass
    break
for par in doc.paragraphs:
    pPr = par._p.pPr
    if pPr is not None and pPr.pStyle is not None and pPr.pStyle.get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "1":
        has_sect = pPr.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr") is not None
        print("H1:", par.text, "| 含分节符:", has_sect)
